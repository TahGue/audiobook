"""Unified document extraction router supporting PDF, EPUB, and DOCX."""

from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import io
import hashlib
import json
import asyncio
import threading
from queue import Queue

router = APIRouter()


class DocumentExtractResponse(BaseModel):
    """Response model for document text extraction."""
    text: str
    total_chars: int
    file_type: str
    file_name: str
    cached: bool = False


def get_file_hash(content: bytes) -> str:
    """Generate MD5 hash of file content for caching."""
    return hashlib.md5(content).hexdigest()


def extract_pdf_text(content: bytes, progress_callback=None) -> str:
    """Extract text from PDF with proper Arabic handling."""
    from services.arabic_text_service import extract_arabic_from_pdf
    
    result = extract_arabic_from_pdf(content, progress_callback=progress_callback)
    return result.text


def extract_epub_text(content: bytes) -> str:
    """Extract text from EPUB using ebooklib."""
    from ebooklib import epub
    from bs4 import BeautifulSoup
    
    # Write to temp file since ebooklib needs a file path
    import tempfile
    import os
    
    with tempfile.NamedTemporaryFile(suffix='.epub', delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    
    try:
        book = epub.read_epub(tmp_path)
        text_parts = []
        
        # Extract text from all HTML documents in the EPUB
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                html_content = item.get_content().decode('utf-8', errors='ignore')
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                text_parts.append(text)
        
        return "\n\n".join(text_parts).strip()
    finally:
        os.unlink(tmp_path)


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    
    doc = Document(io.BytesIO(content))
    text_parts = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n\n".join(text_parts).strip()


@router.post("/extract/")
async def extract_document(
    file: UploadFile = File(...),
    use_cache: bool = True
):
    """
    Extract text from uploaded document (PDF, EPUB, or DOCX) with SSE progress streaming.
    
    Supports:
    - PDF (.pdf)
    - EPUB (.epub)  
    - Word Documents (.docx)
    
    Returns Server-Sent Events stream with progress updates and final result.
    """
    print(f"[DEBUG] File upload: filename={file.filename}, content_type={file.content_type}")
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Read file content
    content = await file.read()
    print(f"[DEBUG] File size: {len(content)} bytes")
    if not content:
        raise HTTPException(status_code=400, detail="Empty file")
    
    # Determine file type
    filename = file.filename.lower()
    
    async def event_generator():
        """Generate SSE events with extraction progress."""
        progress_queue = Queue()
        result_container = {'text': None, 'error': None, 'file_type': None}
        
        def run_extraction():
            """Run extraction in background thread."""
            try:
                if filename.endswith('.pdf'):
                    def progress_callback(progress):
                        progress_queue.put({'type': 'progress', 'progress': progress})
                    
                    text = extract_pdf_text(content, progress_callback=progress_callback)
                    result_container['file_type'] = 'pdf'
                    result_container['text'] = text
                elif filename.endswith('.epub'):
                    progress_queue.put({'type': 'progress', 'progress': 50})
                    text = extract_epub_text(content)
                    result_container['file_type'] = 'epub'
                    result_container['text'] = text
                    progress_queue.put({'type': 'progress', 'progress': 100})
                elif filename.endswith('.docx'):
                    progress_queue.put({'type': 'progress', 'progress': 50})
                    text = extract_docx_text(content)
                    result_container['file_type'] = 'docx'
                    result_container['text'] = text
                    progress_queue.put({'type': 'progress', 'progress': 100})
                else:
                    result_container['error'] = f"Unsupported file type: {file.filename}"
            except Exception as e:
                result_container['error'] = str(e)
        
        # Start extraction in background thread
        extraction_thread = threading.Thread(target=run_extraction)
        extraction_thread.start()
        
        # Stream progress events while extraction runs
        while extraction_thread.is_alive() or not progress_queue.empty():
            while not progress_queue.empty():
                event = progress_queue.get()
                yield f"data: {json.dumps(event)}\n\n"
            await asyncio.sleep(0.1)
        
        # Wait for thread to complete
        extraction_thread.join()
        
        # Send final result or error
        if result_container['error']:
            yield f"data: {json.dumps({'type': 'error', 'detail': result_container['error']})}\n\n"
        elif result_container['text'] is not None:
            text = result_container['text']
            if not text.strip():
                yield f"data: {json.dumps({'type': 'error', 'detail': 'No text could be extracted from the document'})}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'complete', 'text': text, 'total_chars': len(text), 'file_type': result_container['file_type'], 'file_name': file.filename, 'cached': False})}\n\n"
        else:
            yield f"data: {json.dumps({'type': 'error', 'detail': 'Extraction failed unexpectedly'})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.get("/supported-formats/")
async def get_supported_formats():
    """Get list of supported document formats."""
    return {
        "formats": [
            {
                "extension": ".pdf",
                "mime_type": "application/pdf",
                "name": "PDF Document",
                "description": "Portable Document Format"
            },
            {
                "extension": ".epub",
                "mime_type": "application/epub+zip",
                "name": "EPUB eBook",
                "description": "Electronic Publication format"
            },
            {
                "extension": ".docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "name": "Word Document",
                "description": "Microsoft Word Open XML format"
            }
        ],
        "recommended_for_arabic": [".pdf", ".epub"],
        "max_file_size_mb": 50
    }
